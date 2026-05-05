# import streamlit as st
# import re
# import numpy as np
# import pandas as pd
# from io import BytesIO
# import pdfplumber
# import docx
# import spacy

# from sentence_transformers import SentenceTransformer
# from sklearn.metrics.pairwise import cosine_similarity

# # ---------------- PAGE CONFIG ----------------
# st.set_page_config(page_title="AI Resume Screener", layout="wide")

# # ---------------- STYLING ----------------
# st.markdown("""
# <style>
# .stApp {
#     background: linear-gradient(135deg, #0f172a, #020617);
#     color: #e2e8f0;
# }
# .header {
#     font-size: 40px;
#     font-weight: 800;
#     background: linear-gradient(90deg,#38bdf8,#6366f1);
#     -webkit-background-clip: text;
#     -webkit-text-fill-color: transparent;
# }
# .card {
#     background: rgba(255,255,255,0.05);
#     padding: 20px;
#     border-radius: 15px;
#     backdrop-filter: blur(10px);
#     box-shadow: 0 8px 32px rgba(0,0,0,0.3);
# }
# .metric {
#     font-size: 22px;
#     font-weight: bold;
# }
# .green {color:#22c55e;}
# .red {color:#ef4444;}
# </style>
# """, unsafe_allow_html=True)

# # ---------------- LOAD MODELS ----------------
# @st.cache_resource
# def load_models():
#     return SentenceTransformer('all-MiniLM-L6-v2'), spacy.load("en_core_web_sm")

# model, nlp = load_models()

# # ---------------- FILE READ ----------------
# def read_pdf(file):
#     text=""
#     with pdfplumber.open(BytesIO(file.read())) as pdf:
#         for p in pdf.pages:
#             if p.extract_text():
#                 text += p.extract_text()
#     return text

# def read_docx(file):
#     doc = docx.Document(BytesIO(file.read()))
#     return " ".join([p.text for p in doc.paragraphs])

# def read_txt(file):
#     return file.read().decode("utf-8")

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return read_pdf(file)
#     elif file.name.endswith(".docx"):
#         return read_docx(file)
#     elif file.name.endswith(".txt"):
#         return read_txt(file)
#     return ""

# # ---------------- NLP ----------------
# SKILL_MAP={"ml":"machine learning","ai":"artificial intelligence","dl":"deep learning"}
# SKILLS=["python","sql","machine learning","deep learning","numpy","pandas","excel","power bi"]

# def preprocess(text):
#     text=text.lower()
#     text=re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
#     words=text.split()
#     words=[SKILL_MAP.get(w,w) for w in words]
#     return " ".join(words)

# def extract_skills(text):
#     return list(set([s for s in SKILLS if s in text]))

# def extract_experience(text):
#     matches=re.findall(r'(\d+)\+?\s*(years|yrs)', text)
#     return max([int(m[0]) for m in matches]) if matches else 0

# def extract_education(text):
#     deg=["bachelor","master","btech","mtech","bsc","msc"]
#     return [d for d in deg if d in text]

# def extract_certifications(text):
#     platforms=["coursera","udemy","ibm","google","aws","azure","kaggle","cognitive class"]
#     return [p for p in platforms if p in text]

# def project_score(text):
#     return min(sum([1 for k in ["project","developed","built","implemented"] if k in text])/3,1)

# # ---------------- SIDEBAR ----------------
# st.sidebar.title("⚙️ Controls")
# uploaded_files = st.sidebar.file_uploader("Upload Resumes", accept_multiple_files=True)
# job_desc = st.sidebar.text_area("Job Description")

# analyze = st.sidebar.button("Analyze")

# # ---------------- HEADER ----------------
# st.markdown('<div class="header">🚀 AI Resume Screening Dashboard</div>', unsafe_allow_html=True)
# st.markdown("Smart hiring powered by AI & NLP")

# # ---------------- PROCESS ----------------
# if analyze:

#     if not uploaded_files or not job_desc:
#         st.warning("Upload resumes and enter job description")
#         st.stop()

#     job_clean = preprocess(job_desc)
#     jd_skills = extract_skills(job_clean)
#     jd_exp = extract_experience(job_clean)
#     jd_edu = extract_education(job_clean)

#     job_emb = model.encode([job_clean])[0]

#     results=[]

#     for file in uploaded_files:
#         text = extract_text(file)
#         clean = preprocess(text)

#         emb=model.encode([clean])[0]
#         sim=cosine_similarity([emb],[job_emb])[0][0]

#         skills=extract_skills(clean)
#         exp=extract_experience(clean)
#         edu=extract_education(clean)
#         cert=extract_certifications(clean)

#         skill_sc=len(set(skills)&set(jd_skills))/len(jd_skills) if jd_skills else 0
#         exp_sc=min(exp/jd_exp,1) if jd_exp else 1
#         edu_sc=1 if set(edu)&set(jd_edu) else 0 if jd_edu else 1
#         cert_sc=min(len(cert)/3,1)
#         proj_sc=project_score(clean)

#         final_score=(0.5*sim+0.2*skill_sc+0.15*exp_sc+0.1*edu_sc+0.05*cert_sc+0.05*proj_sc)

#         decision="Suitable" if final_score>=0.5 else "Not Suitable"

#         results.append({
#             "Name":file.name,
#             "Score":round(final_score,3),
#             "Decision":decision,
#             "Skills":skills,
#             "Experience":exp,
#             "Education":edu,
#             "Certifications":cert
#         })

#     df=pd.DataFrame(results).sort_values(by="Score",ascending=False)

#     # ---------------- METRICS ----------------
#     col1,col2,col3=st.columns(3)

#     col1.markdown(f'<div class="card"><div class="metric">Candidates</div>{len(df)}</div>',unsafe_allow_html=True)
#     col2.markdown(f'<div class="card"><div class="metric">Selected</div>{len(df[df["Decision"]=="Suitable"])}</div>',unsafe_allow_html=True)
#     col3.markdown(f'<div class="card"><div class="metric">Rejected</div>{len(df[df["Decision"]=="Not Suitable"])}</div>',unsafe_allow_html=True)

#     # ---------------- TOP CANDIDATE ----------------
#     top=df.iloc[0]
#     color="green" if top["Decision"]=="Suitable" else "red"

#     st.markdown("## 🏆 Top Candidate")
#     st.markdown(f"""
#     <div class="card">
#     <h3>{top['Name']}</h3>
#     <p>Score: {top['Score']}</p>
#     <p class="{color}">{top['Decision']}</p>
#     </div>
#     """,unsafe_allow_html=True)

#     # ---------------- TABLE ----------------
#     st.markdown("## 📊 Ranking Table")
#     st.dataframe(df,use_container_width=True)

#     # ---------------- CHART ----------------
#     st.markdown("## 📈 Score Chart")
#     st.bar_chart(df.set_index("Name")["Score"])

#     # ---------------- INSIGHTS ----------------
#     st.markdown("## 🔍 Candidate Profiles")

#     for _,row in df.iterrows():
#         color="green" if row["Decision"]=="Suitable" else "red"

#         with st.expander(row["Name"]):
#             st.markdown(f"""
#             <div class="card">
#             <p><b>Score:</b> {row['Score']}</p>
#             <p class="{color}">{row['Decision']}</p>
#             <p><b>Skills:</b> {row['Skills']}</p>
#             <p><b>Experience:</b> {row['Experience']} years</p>
#             <p><b>Education:</b> {row['Education']}</p>
#             <p><b>Certifications:</b> {row['Certifications']}</p>
#             </div>
#             """,unsafe_allow_html=True)

#     # ---------------- DOWNLOAD ----------------
#     csv=df.to_csv(index=False).encode()
#     st.download_button("⬇ Download Results",csv,"results.csv")




# import streamlit as st
# import re
# import pandas as pd
# from io import BytesIO
# import pdfplumber
# import docx
# from sentence_transformers import SentenceTransformer
# from sklearn.metrics.pairwise import cosine_similarity

# # ---------------- PAGE ----------------
# st.set_page_config(layout="wide")

# # ---------------- STYLE ----------------
# st.markdown("""
# <style>
# .stApp {background-color:#0f172a;}
# .title {font-size:32px;font-weight:700;color:white;}
# .subtitle {color:#94a3b8;margin-bottom:25px;}
# .card {
#     background:#1e293b;
#     padding:18px;
#     border-radius:10px;
#     border:1px solid #334155;
#     margin-bottom:12px;
# }
# .tag {
#     background:#334155;
#     padding:4px 8px;
#     border-radius:6px;
#     margin:2px;
#     display:inline-block;
#     font-size:12px;
# }
# .green {color:#22c55e;}
# .red {color:#ef4444;}
# </style>
# """, unsafe_allow_html=True)

# # ---------------- MODEL ----------------
# @st.cache_resource
# def load_model():
#     return SentenceTransformer('all-MiniLM-L6-v2')

# model = load_model()

# # ---------------- FILE READ ----------------
# def read_pdf(file):
#     text=""
#     with pdfplumber.open(BytesIO(file.read())) as pdf:
#         for p in pdf.pages:
#             if p.extract_text():
#                 text += p.extract_text()
#     return text

# def read_docx(file):
#     doc = docx.Document(BytesIO(file.read()))
#     return " ".join([p.text for p in doc.paragraphs])

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return read_pdf(file)
#     elif file.name.endswith(".docx"):
#         return read_docx(file)
#     else:
#         return file.read().decode("utf-8")

# # ---------------- NLP ----------------
# SKILL_MAP = {"ml":"machine learning","ai":"artificial intelligence","dl":"deep learning"}

# SKILLS = ["python","sql","machine learning","deep learning","numpy","pandas","excel","power bi"]

# def preprocess(text):
#     text = text.lower()
#     text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
#     words = text.split()
#     words = [SKILL_MAP.get(w, w) for w in words]
#     return " ".join(words)

# def extract_skills(text):
#     return list(set([s for s in SKILLS if s in text]))

# def extract_experience(text):
#     matches = re.findall(r'(\d+)\+?\s*(years|yrs)', text)
#     return max([int(m[0]) for m in matches]) if matches else 0

# def extract_education(text):
#     deg = ["bachelor","master","btech","mtech","bsc","msc"]
#     return [d for d in deg if d in text]

# def extract_certifications(text):
#     platforms = ["coursera","udemy","ibm","google","aws","azure","kaggle","cognitive class"]
#     return [p for p in platforms if p in text]

# def project_score(text):
#     return min(sum([1 for k in ["project","developed","built","implemented"] if k in text])/3,1)

# # ---------------- HEADER ----------------
# st.markdown('<div class="title">AI Resume Screening System</div>', unsafe_allow_html=True)
# st.markdown('<div class="subtitle">Full NLP Pipeline + Intelligent Candidate Ranking</div>', unsafe_allow_html=True)

# # ---------------- INPUT ----------------
# col1, col2 = st.columns(2)

# with col1:
#     files = st.file_uploader("Upload Resumes", accept_multiple_files=True)

# with col2:
#     job_desc = st.text_area("Job Description", height=150)

# analyze = st.button("Analyze Candidates")

# # ---------------- PROCESS ----------------
# if analyze:

#     job_clean = preprocess(job_desc)

#     jd_skills = extract_skills(job_clean)
#     jd_exp = extract_experience(job_clean)
#     jd_edu = extract_education(job_clean)

#     job_emb = model.encode([job_clean])[0]

#     results = []

#     for file in files:
#         text = extract_text(file)
#         clean = preprocess(text)

#         emb = model.encode([clean])[0]
#         sim = cosine_similarity([emb], [job_emb])[0][0]

#         skills = extract_skills(clean)
#         exp = extract_experience(clean)
#         edu = extract_education(clean)
#         cert = extract_certifications(clean)

#         skill_sc = len(set(skills)&set(jd_skills))/len(jd_skills) if jd_skills else 0
#         exp_sc = min(exp/jd_exp,1) if jd_exp else 1
#         edu_sc = 1 if set(edu)&set(jd_edu) else 0 if jd_edu else 1
#         cert_sc = min(len(cert)/3,1)
#         proj_sc = project_score(clean)

#         final_score = (
#             0.5*sim +
#             0.2*skill_sc +
#             0.15*exp_sc +
#             0.1*edu_sc +
#             0.05*cert_sc +
#             0.05*proj_sc
#         )

#         decision = "Suitable" if final_score >= 0.5 else "Not Suitable"

#         results.append({
#             "Name": file.name,
#             "Score": round(final_score,3),
#             "Decision": decision,
#             "Skills": skills,
#             "Experience": exp,
#             "Education": edu,
#             "Certifications": cert
#         })

#     df = pd.DataFrame(results).sort_values(by="Score", ascending=False)

#     # ---------------- METRICS ----------------
#     c1,c2,c3 = st.columns(3)
#     c1.markdown(f'<div class="card"><b>Total</b><br>{len(df)}</div>', unsafe_allow_html=True)
#     c2.markdown(f'<div class="card"><b>Selected</b><br>{len(df[df["Decision"]=="Suitable"])}</div>', unsafe_allow_html=True)
#     c3.markdown(f'<div class="card"><b>Rejected</b><br>{len(df[df["Decision"]=="Not Suitable"])}</div>', unsafe_allow_html=True)

#     # ---------------- TOP ----------------
#     top = df.iloc[0]
#     color = "green" if top["Decision"]=="Suitable" else "red"

#     st.markdown("### Top Candidate")
#     st.markdown(f"""
#     <div class="card">
#     <b>{top['Name']}</b><br>
#     Score: {top['Score']}<br>
#     <span class="{color}">{top['Decision']}</span>
#     </div>
#     """, unsafe_allow_html=True)

#     # ---------------- TABLE ----------------
#     st.markdown("### Ranking")
#     st.dataframe(df, use_container_width=True)

#     # ---------------- DETAILS ----------------
#     st.markdown("### Candidate Details")

#     for _, row in df.iterrows():
#         color = "green" if row["Decision"]=="Suitable" else "red"

#         st.markdown(f"""
#         <div class="card">
#         <b>{row['Name']}</b><br>
#         Score: {row['Score']}<br>
#         <span class="{color}">{row['Decision']}</span><br><br>

#         <b>Skills:</b><br>
#         {" ".join([f"<span class='tag'>{s}</span>" for s in row["Skills"]])}<br><br>

#         <b>Experience:</b> {row['Experience']} years<br>
#         <b>Education:</b> {row['Education']}<br>
#         <b>Certifications:</b> {row['Certifications']}
#         </div>
#         """, unsafe_allow_html=True)

#     # ---------------- CHART ----------------
#     st.bar_chart(df.set_index("Name")["Score"])

#     # ---------------- DOWNLOAD ----------------
#     csv = df.to_csv(index=False).encode()
#     st.download_button("Download Results", csv, "results.csv")




import streamlit as st
import re
import pandas as pd
from io import BytesIO
import pdfplumber
import docx
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# ---------------- PAGE ----------------
st.set_page_config(layout="wide")

# ---------------- STYLE ----------------
st.markdown("""
<style>
.stApp {
    background-color: #0f172a;
}

/* Center container */
.main-container {
    max-width: 900px;
    margin: auto;
}

/* Title */
.title {
    font-size: 34px;
    font-weight: 700;
    color: white;
    text-align: center;
    margin-bottom: 5px;
}

/* Subtitle */
.subtitle {
    text-align: center;
    color: #94a3b8;
    margin-bottom: 30px;
}

/* Cards */
.card {
    background: #1e293b;
    padding: 22px;
    border-radius: 12px;
    border: 1px solid #334155;
    margin-bottom: 18px;
}

/* Tags */
.tag {
    background: #334155;
    padding: 6px 10px;
    border-radius: 8px;
    margin: 3px;
    display: inline-block;
    font-size: 12px;
}

/* Colors */
.green {color:#22c55e;}
.red {color:#ef4444;}
</style>
""", unsafe_allow_html=True)

# ---------------- MODEL ----------------
@st.cache_resource
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

model = load_model()

# ---------------- FILE READ ----------------
def read_pdf(file):
    text=""
    with pdfplumber.open(BytesIO(file.read())) as pdf:
        for p in pdf.pages:
            if p.extract_text():
                text += p.extract_text()
    return text

def read_docx(file):
    doc = docx.Document(BytesIO(file.read()))
    return " ".join([p.text for p in doc.paragraphs])

def extract_text(file):
    if file.name.endswith(".pdf"):
        return read_pdf(file)
    elif file.name.endswith(".docx"):
        return read_docx(file)
    else:
        return file.read().decode("utf-8")

# ---------------- NLP ----------------
SKILL_MAP = {"ml":"machine learning","ai":"artificial intelligence","dl":"deep learning"}

SKILLS = ["python","sql","machine learning","deep learning","numpy","pandas","excel","power bi"]

def preprocess(text):
    text = text.lower()
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    words = text.split()
    words = [SKILL_MAP.get(w, w) for w in words]
    return " ".join(words)

def extract_skills(text):
    return list(set([s for s in SKILLS if s in text]))

def extract_experience(text):
    matches = re.findall(r'(\d+)\+?\s*(years|yrs)', text)
    return max([int(m[0]) for m in matches]) if matches else 0

def extract_education(text):
    deg = ["bachelor","master","btech","mtech","bsc","msc"]
    return [d for d in deg if d in text]

def extract_certifications(text):
    platforms = ["coursera","udemy","ibm","google","aws","azure","kaggle","cognitive class"]
    return [p for p in platforms if p in text]

def project_score(text):
    return min(sum([1 for k in ["project","developed","built","implemented"] if k in text])/3,1)

# ---------------- CONTAINER START ----------------
st.markdown('<div class="main-container">', unsafe_allow_html=True)

# ---------------- HEADER ----------------
st.markdown('<div class="title">AI Resume Screening System</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Smart Candidate Matching using NLP + AI</div>', unsafe_allow_html=True)

# ---------------- INPUT ----------------
st.markdown("### Upload Resumes")
files = st.file_uploader("Upload multiple resumes", accept_multiple_files=True)

st.markdown("### Job Description")
job_desc = st.text_area("Paste job description", height=200)

st.markdown("")
analyze = st.button("🔍 Analyze Candidates")

# ---------------- PROCESS ----------------
if analyze:

    job_clean = preprocess(job_desc)

    jd_skills = extract_skills(job_clean)
    jd_exp = extract_experience(job_clean)
    jd_edu = extract_education(job_clean)

    job_emb = model.encode([job_clean])[0]

    results = []

    for file in files:
        text = extract_text(file)
        clean = preprocess(text)

        emb = model.encode([clean])[0]
        sim = cosine_similarity([emb], [job_emb])[0][0]

        skills = extract_skills(clean)
        exp = extract_experience(clean)
        edu = extract_education(clean)
        cert = extract_certifications(clean)

        skill_sc = len(set(skills)&set(jd_skills))/len(jd_skills) if jd_skills else 0
        exp_sc = min(exp/jd_exp,1) if jd_exp else 1
        edu_sc = 1 if set(edu)&set(jd_edu) else 0 if jd_edu else 1
        cert_sc = min(len(cert)/3,1)
        proj_sc = project_score(clean)

        final_score = (
            0.5*sim +
            0.2*skill_sc +
            0.15*exp_sc +
            0.1*edu_sc +
            0.05*cert_sc +
            0.05*proj_sc
        )

        decision = "Suitable" if final_score >= 0.5 else "Not Suitable"

        results.append({
            "Name": file.name,
            "Score": round(final_score,3),
            "Decision": decision,
            "Skills": skills,
            "Experience": exp,
            "Education": edu,
            "Certifications": cert
        })

    df = pd.DataFrame(results).sort_values(by="Score", ascending=False)

    # ---------------- SUMMARY ----------------
    st.markdown("### Summary")

    st.markdown(f"""
    <div class="card">
    Total Candidates: {len(df)}<br>
    Selected: {len(df[df["Decision"]=="Suitable"])}<br>
    Rejected: {len(df[df["Decision"]=="Not Suitable"])}
    </div>
    """, unsafe_allow_html=True)

    # ---------------- TOP ----------------
    top = df.iloc[0]
    color = "green" if top["Decision"]=="Suitable" else "red"

    st.markdown("### Top Candidate")
    st.markdown(f"""
    <div class="card">
    <b>{top['Name']}</b><br><br>
    Score: {top['Score']}<br>
    <span class="{color}">{top['Decision']}</span>
    </div>
    """, unsafe_allow_html=True)

    # ---------------- TABLE ----------------
    st.markdown("### Ranking")
    st.dataframe(df, use_container_width=True)

    # ---------------- DETAILS ----------------
    st.markdown("### Candidate Details")

    for _, row in df.iterrows():
        color = "green" if row["Decision"]=="Suitable" else "red"

        st.markdown(f"""
        <div class="card">
        <b>{row['Name']}</b><br><br>
        Score: {row['Score']}<br>
        <span class="{color}">{row['Decision']}</span><br><br>

        <b>Skills:</b><br>
        {" ".join([f"<span class='tag'>{s}</span>" for s in row["Skills"]])}<br><br>

        <b>Experience:</b> {row['Experience']} years<br>
        <b>Education:</b> {row['Education']}<br>
        <b>Certifications:</b> {row['Certifications']}
        </div>
        """, unsafe_allow_html=True)

    # ---------------- CHART ----------------
    st.bar_chart(df.set_index("Name")["Score"])

    # ---------------- DOWNLOAD ----------------
    csv = df.to_csv(index=False).encode()
    st.download_button("⬇ Download Results", csv, "results.csv")

# ---------------- CONTAINER END ----------------
st.markdown('</div>', unsafe_allow_html=True)
